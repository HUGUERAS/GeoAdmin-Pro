from __future__ import annotations

from datetime import datetime, timezone
import io
import json
from pathlib import Path
from typing import Any
from uuid import uuid4
import zipfile

from pyproj import Transformer
from shapely.geometry import Polygon
from shapely.ops import transform

from integracoes.referencia_cliente import resumir_vertices

try:
    from docx import Document
except Exception:  # pragma: no cover - fallback defensivo
    Document = None


BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
UPLOADS_DIR = DATA_DIR / "formulario_uploads"
TEMPLATES_DIR = BASE_DIR / "static" / "templates"
TEMPLATE_CARTA_CONFRONTACAO = TEMPLATES_DIR / "carta_confrontacao.docx"


def _get_supabase():
    from main import get_supabase
    return get_supabase()


def _agora_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _data_extenso_ptbr(data: datetime | None = None) -> str:
    meses = [
        "janeiro",
        "fevereiro",
        "março",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    ]
    valor = data or datetime.now()
    return f"{valor.day} de {meses[valor.month - 1]} de {valor.year}"


def _vertices_validos(vertices: list[dict[str, Any]] | None) -> list[dict[str, float]]:
    lista = vertices or []
    if len(lista) < 3:
        return []
    normalizados: list[dict[str, float]] = []
    for item in lista:
        normalizados.append({"lon": float(item["lon"]), "lat": float(item["lat"])})
    return normalizados


def _polygon_from_vertices(vertices: list[dict[str, Any]]) -> Polygon | None:
    coords = [(float(item["lon"]), float(item["lat"])) for item in vertices]
    if len(coords) < 3:
        return None
    polygon = Polygon(coords)
    if not polygon.is_valid:
        polygon = polygon.buffer(0)
    if polygon.is_empty:
        return None
    return polygon


def _utm_epsg(lat: float, lon: float) -> int:
    fuso = int((lon + 180) // 6) + 1
    return 32700 + fuso if lat < 0 else 32600 + fuso


def _transformar_para_metros(polygon: Polygon) -> Polygon:
    centroide = polygon.centroid
    epsg = _utm_epsg(float(centroide.y), float(centroide.x))
    transformer = Transformer.from_crs("EPSG:4326", f"EPSG:{epsg}", always_xy=True)
    return transform(transformer.transform, polygon)


def _resumo_vertices(vertices: list[dict[str, Any]] | None) -> dict[str, Any] | None:
    vertices_ok = _vertices_validos(vertices)
    if not vertices_ok:
        return None
    return resumir_vertices(vertices_ok)


def _geometria_preferencial(area: dict[str, Any]) -> tuple[str, list[dict[str, Any]]]:
    geometria_final = _vertices_validos(area.get("geometria_final"))
    if geometria_final:
        return "final", geometria_final
    esboco = _vertices_validos(area.get("geometria_esboco"))
    return "esboco", esboco


def _status_geometria(area: dict[str, Any]) -> str:
    if _vertices_validos(area.get("geometria_final")):
        return "geometria_final"
    if _vertices_validos(area.get("geometria_esboco")):
        return "apenas_esboco"
    return "sem_geometria"


def _row_para_area(raw: dict[str, Any] | None) -> dict[str, Any] | None:
    if not raw:
        return None
    return {
        "id": raw.get("id"),
        "projeto_id": raw.get("projeto_id"),
        "cliente_id": raw.get("cliente_id"),
        "nome": raw.get("nome"),
        "proprietario_nome": raw.get("proprietario_nome"),
        "municipio": raw.get("municipio"),
        "estado": raw.get("estado"),
        "comarca": raw.get("comarca"),
        "matricula": raw.get("matricula"),
        "ccir": raw.get("ccir"),
        "car": raw.get("car"),
        "observacoes": raw.get("observacoes"),
        "origem_tipo": raw.get("origem_tipo") or "manual",
        "geometria_esboco": raw.get("geometria_esboco") or [],
        "geometria_final": raw.get("geometria_final") or [],
        "resumo_esboco": raw.get("resumo_esboco"),
        "resumo_final": raw.get("resumo_final"),
        "anexos": raw.get("anexos") or [],
        "criado_em": raw.get("criado_em") or raw.get("created_at"),
        "atualizado_em": raw.get("atualizado_em") or raw.get("updated_at"),
        "deleted_at": raw.get("deleted_at"),
    }


def _normalizar_area(area: dict[str, Any]) -> dict[str, Any]:
    tipo_geometria, geometria_ativa = _geometria_preferencial(area)
    resumo_esboco = area.get("resumo_esboco") or _resumo_vertices(area.get("geometria_esboco"))
    resumo_final = area.get("resumo_final") or _resumo_vertices(area.get("geometria_final"))
    resumo_ativo = resumo_final if tipo_geometria == "final" else resumo_esboco

    return {
        **area,
        "status_geometria": _status_geometria(area),
        "tipo_geometria_ativa": tipo_geometria if geometria_ativa else None,
        "geometria_ativa": geometria_ativa,
        "resumo_esboco": resumo_esboco,
        "resumo_final": resumo_final,
        "resumo_ativo": resumo_ativo,
        "anexos": area.get("anexos") or [],
    }


def _substituir_placeholders_texto(texto: str, campos: dict[str, Any]) -> str:
    resultado = texto
    for chave, valor in campos.items():
        resultado = resultado.replace(f"{{{{ {chave} }}}}", str(valor))
        resultado = resultado.replace(f"{{{{{chave}}}}}", str(valor))
    return resultado


def _substituir_placeholders_docx(documento, campos: dict[str, Any]) -> None:
    def atualizar_paragrafo(paragrafo) -> None:
        texto = _substituir_placeholders_texto(paragrafo.text, campos)
        if texto == paragrafo.text:
            return
        if paragrafo.runs:
            paragrafo.runs[0].text = texto
            for run in paragrafo.runs[1:]:
                run.text = ""
        else:
            paragrafo.add_run(texto)

    for paragrafo in documento.paragraphs:
        atualizar_paragrafo(paragrafo)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    atualizar_paragrafo(paragrafo)


def _gerar_carta_confrontacao_docx(
    *,
    projeto: dict[str, Any],
    area_a: dict[str, Any],
    area_b: dict[str, Any],
    confronto: dict[str, Any],
) -> bytes | None:
    if Document is None or not TEMPLATE_CARTA_CONFRONTACAO.exists():
        return None

    documento = Document(str(TEMPLATE_CARTA_CONFRONTACAO))
    campos = {
        "projeto_nome": projeto.get("projeto_nome") or projeto.get("nome") or "Projeto",
        "municipio": projeto.get("municipio") or area_a.get("municipio") or area_b.get("municipio") or "",
        "comarca": projeto.get("comarca") or area_a.get("comarca") or area_b.get("comarca") or "",
        "data_extenso": _data_extenso_ptbr(),
        "proprietario_a": area_a.get("proprietario_nome") or "Nao informado",
        "cpf_a": area_a.get("cpf") or area_a.get("proprietario_cpf") or "",
        "imovel_a": area_a.get("nome") or "Area sem nome",
        "matricula_a": area_a.get("matricula") or "",
        "proprietario_b": area_b.get("proprietario_nome") or "Nao informado",
        "cpf_b": area_b.get("cpf") or area_b.get("proprietario_cpf") or "",
        "imovel_b": area_b.get("nome") or "Area sem nome",
        "matricula_b": area_b.get("matricula") or "",
        "tipo_relacao": confronto.get("tipo") or "divisa",
        "comprimento_contato": f"{confronto.get('contato_m', 0)} m",
        "area_sobreposicao": f"{confronto.get('area_intersecao_ha', 0)} ha",
    }
    _substituir_placeholders_docx(documento, campos)

    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def listar_areas_projeto(projeto_id: str, sb=None) -> list[dict[str, Any]]:
    cliente = sb or _get_supabase()
    res = (
        cliente.table("areas_projeto")
        .select("*")
        .eq("projeto_id", projeto_id)
        .is_("deleted_at", "null")
        .order("atualizado_em", desc=True)
        .execute()
    )
    areas = [_normalizar_area(_row_para_area(area) or {}) for area in (res.data or [])]
    return areas


def obter_area(area_id: str, sb=None) -> dict[str, Any] | None:
    cliente = sb or _get_supabase()
    res = (
        cliente.table("areas_projeto")
        .select("*")
        .eq("id", area_id)
        .is_("deleted_at", "null")
        .maybe_single()
        .execute()
    )
    area = _row_para_area(res.data)
    if not area:
        return None
    return _normalizar_area(area)


def salvar_area_projeto(
    *,
    projeto_id: str,
    cliente_id: str | None,
    nome: str,
    proprietario_nome: str | None = None,
    municipio: str | None = None,
    estado: str | None = None,
    comarca: str | None = None,
    matricula: str | None = None,
    ccir: str | None = None,
    car: str | None = None,
    observacoes: str | None = None,
    origem_tipo: str = "formulario",
    geometria_esboco: list[dict[str, Any]] | None = None,
    geometria_final: list[dict[str, Any]] | None = None,
    anexos: list[dict[str, Any]] | None = None,
    area_id: str | None = None,
    sb=None,
) -> dict[str, Any]:
    cliente = sb or _get_supabase()
    existente = obter_area(area_id, sb=cliente) if area_id else None

    vertices_esboco = (
        _vertices_validos(geometria_esboco)
        if geometria_esboco is not None
        else _vertices_validos((existente or {}).get("geometria_esboco"))
    )
    vertices_final = (
        _vertices_validos(geometria_final)
        if geometria_final is not None
        else _vertices_validos((existente or {}).get("geometria_final"))
    )

    payload = {
        "id": area_id or (existente or {}).get("id") or str(uuid4()),
        "projeto_id": projeto_id,
        "cliente_id": cliente_id,
        "nome": (nome or "").strip() or "Area sem nome",
        "proprietario_nome": proprietario_nome,
        "municipio": municipio,
        "estado": estado,
        "comarca": comarca,
        "matricula": matricula,
        "ccir": ccir,
        "car": car,
        "observacoes": observacoes,
        "origem_tipo": origem_tipo,
        "geometria_esboco": vertices_esboco,
        "geometria_final": vertices_final,
        "resumo_esboco": _resumo_vertices(vertices_esboco),
        "resumo_final": _resumo_vertices(vertices_final),
        "anexos": anexos if anexos is not None else ((existente or {}).get("anexos") or []),
        "deleted_at": None,
    }

    if existente:
        res = cliente.table("areas_projeto").update(payload).eq("id", payload["id"]).execute()
    else:
        res = cliente.table("areas_projeto").insert(payload).execute()

    registro = None
    if res.data:
        registro = _row_para_area(res.data[0])
    if not registro:
        registro = _row_para_area(
            cliente.table("areas_projeto").select("*").eq("id", payload["id"]).maybe_single().execute().data
        )
    if not registro:
        raise RuntimeError("Falha ao persistir area do projeto no Supabase.")
    return _normalizar_area(registro)


def anexar_arquivos_area(
    *,
    area_id: str,
    cliente_id: str | None,
    arquivos: list[tuple[str, bytes, str | None]],
    sb=None,
) -> list[dict[str, Any]]:
    if not arquivos:
        return []

    cliente = sb or _get_supabase()
    area = obter_area(area_id, sb=cliente)
    if not area:
        return []

    pasta_cliente = UPLOADS_DIR / (cliente_id or "sem-cliente") / area_id
    pasta_cliente.mkdir(parents=True, exist_ok=True)
    anexos = list(area.get("anexos") or [])

    for nome_original, conteudo, content_type in arquivos:
        extensao = Path(nome_original or "arquivo.bin").suffix or ".bin"
        nome_seguro = f"{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}{extensao}"
        destino = pasta_cliente / nome_seguro
        destino.write_bytes(conteudo)
        anexos.append(
            {
                "id": str(uuid4()),
                "nome_original": nome_original,
                "arquivo_nome": nome_seguro,
                "content_type": content_type,
                "tamanho_bytes": len(conteudo),
                "caminho_local": str(destino),
                "criado_em": _agora_iso(),
            }
        )

    cliente.table("areas_projeto").update({"anexos": anexos}).eq("id", area_id).execute()
    return anexos


def sintetizar_areas_do_projeto(
    *,
    projeto: dict[str, Any],
    cliente: dict[str, Any] | None,
    perimetro_ativo: dict[str, Any] | None,
    geometria_referencia: dict[str, Any] | None,
    sb=None,
) -> list[dict[str, Any]]:
    areas = listar_areas_projeto(projeto.get("id"), sb=sb)
    if areas:
        return areas

    sinteticas: list[dict[str, Any]] = []
    projeto_id = projeto.get("id")
    cliente_id = cliente.get("id") if cliente else projeto.get("cliente_id")
    proprietario = (cliente or {}).get("nome") or projeto.get("cliente_nome")
    nome_base = projeto.get("nome_imovel") or projeto.get("projeto_nome") or "Area principal"

    if geometria_referencia and _vertices_validos(geometria_referencia.get("vertices")):
        sinteticas.append(
            _normalizar_area(
                {
                    "id": f"{projeto_id}-ref",
                    "projeto_id": projeto_id,
                    "cliente_id": cliente_id,
                    "nome": f"{nome_base} - Esboco",
                    "proprietario_nome": proprietario,
                    "municipio": projeto.get("municipio"),
                    "estado": projeto.get("estado"),
                    "comarca": projeto.get("comarca"),
                    "matricula": projeto.get("matricula"),
                    "origem_tipo": geometria_referencia.get("origem_tipo") or "referencia_cliente",
                    "geometria_esboco": geometria_referencia.get("vertices") or [],
                    "geometria_final": [],
                    "anexos": geometria_referencia.get("anexos") or [],
                    "criado_em": geometria_referencia.get("atualizado_em"),
                    "atualizado_em": geometria_referencia.get("atualizado_em"),
                }
            )
        )

    if perimetro_ativo and _vertices_validos(perimetro_ativo.get("vertices")):
        sinteticas.append(
            _normalizar_area(
                {
                    "id": f"{projeto_id}-tec",
                    "projeto_id": projeto_id,
                    "cliente_id": cliente_id,
                    "nome": f"{nome_base} - Tecnico",
                    "proprietario_nome": proprietario,
                    "municipio": projeto.get("municipio"),
                    "estado": projeto.get("estado"),
                    "comarca": projeto.get("comarca"),
                    "matricula": projeto.get("matricula"),
                    "origem_tipo": "perimetro_tecnico",
                    "geometria_esboco": [],
                    "geometria_final": perimetro_ativo.get("vertices") or [],
                    "anexos": [],
                    "criado_em": perimetro_ativo.get("criado_em"),
                    "atualizado_em": perimetro_ativo.get("criado_em"),
                }
            )
        )

    return sinteticas


def detectar_confrontacoes(areas: list[dict[str, Any]]) -> list[dict[str, Any]]:
    confrontacoes: list[dict[str, Any]] = []
    areas_validas: list[dict[str, Any]] = []
    for area in areas:
        _, vertices = _geometria_preferencial(area)
        polygon = _polygon_from_vertices(vertices)
        if polygon is None:
            continue
        areas_validas.append({**area, "_polygon": polygon})

    for indice, area_a in enumerate(areas_validas):
        for area_b in areas_validas[indice + 1:]:
            poly_a = area_a["_polygon"]
            poly_b = area_b["_polygon"]
            if not poly_a.intersects(poly_b) and not poly_a.touches(poly_b):
                continue

            poly_a_m = _transformar_para_metros(poly_a)
            poly_b_m = _transformar_para_metros(poly_b)
            inter = poly_a_m.intersection(poly_b_m)
            contato = poly_a_m.boundary.intersection(poly_b_m.boundary)

            area_intersecao = float(inter.area) if not inter.is_empty else 0.0
            contato_m = float(contato.length) if not contato.is_empty else 0.0
            tipo = "sobreposicao" if area_intersecao > 0.01 else "divisa"

            confrontacoes.append(
                {
                    "id": f"{area_a['id']}::{area_b['id']}",
                    "tipo": tipo,
                    "status": "detectada",
                    "origem": "geometria",
                    "area_a": {
                        "id": area_a.get("id"),
                        "nome": area_a.get("nome"),
                        "proprietario_nome": area_a.get("proprietario_nome"),
                    },
                    "area_b": {
                        "id": area_b.get("id"),
                        "nome": area_b.get("nome"),
                        "proprietario_nome": area_b.get("proprietario_nome"),
                    },
                    "contato_m": round(contato_m, 2),
                    "area_intersecao_ha": round(area_intersecao / 10000, 4),
                }
            )

    confrontacoes.sort(key=lambda item: (item["tipo"] != "sobreposicao", item["area_a"]["nome"] or ""))
    return confrontacoes


def gerar_cartas_confrontacao_zip(
    *,
    projeto: dict[str, Any],
    areas: list[dict[str, Any]],
    confrontacoes: list[dict[str, Any]],
) -> bytes:
    mapa_areas = {area.get("id"): area for area in areas}
    zip_buffer = io.BytesIO()
    nome_projeto = projeto.get("projeto_nome") or projeto.get("nome") or "Projeto"

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        if not confrontacoes:
            zf.writestr(
                "LEIA-ME.txt",
                (
                    f"Nenhuma confrontacao geometrica foi detectada para {nome_projeto}.\n"
                    "Cadastre mais areas ou confirme confrontantes manualmente no app.\n"
                ).encode("utf-8"),
            )
        for indice, confronto in enumerate(confrontacoes, start=1):
            area_a = mapa_areas.get(confronto["area_a"]["id"], confronto["area_a"])
            area_b = mapa_areas.get(confronto["area_b"]["id"], confronto["area_b"])
            docx_bytes = _gerar_carta_confrontacao_docx(
                projeto=projeto,
                area_a=area_a,
                area_b=area_b,
                confronto=confronto,
            )
            if docx_bytes:
                nome_docx = f"CARTA_CONFRONTACAO_{indice:02d}.docx"
                zf.writestr(nome_docx, docx_bytes)
                continue

            texto = f"""
CARTA DE CONFRONTACAO
=====================
Projeto: {nome_projeto}

Area principal:
- Nome: {area_a.get('nome') or 'Area sem nome'}
- Proprietario: {area_a.get('proprietario_nome') or 'Nao informado'}
- Matricula: {area_a.get('matricula') or 'Nao informada'}

Area confrontante:
- Nome: {area_b.get('nome') or 'Area sem nome'}
- Proprietario: {area_b.get('proprietario_nome') or 'Nao informado'}
- Matricula: {area_b.get('matricula') or 'Nao informada'}

Resumo tecnico:
- Tipo de relacao: {confronto.get('tipo')}
- Comprimento aproximado de contato: {confronto.get('contato_m', 0)} m
- Sobreposicao aproximada: {confronto.get('area_intersecao_ha', 0)} ha

Observacao:
Esta carta foi preparada automaticamente a partir das geometrias disponiveis
no GeoAdmin. Revise os dados pessoais, o esboco do cliente e o perimetro
tecnico antes da assinatura definitiva.
""".strip()
            nome_txt = f"CARTA_CONFRONTACAO_{indice:02d}.txt"
            zf.writestr(nome_txt, texto.encode("utf-8"))

        zf.writestr(
            "manifesto_cartas.json",
            json.dumps(
                {
                    "gerado_em": _agora_iso(),
                    "projeto": {
                        "id": projeto.get("id"),
                        "nome": nome_projeto,
                    },
                    "areas_total": len(areas),
                    "confrontacoes_total": len(confrontacoes),
                },
                ensure_ascii=False,
                indent=2,
            ).encode("utf-8"),
        )

    zip_buffer.seek(0)
    return zip_buffer.getvalue()
